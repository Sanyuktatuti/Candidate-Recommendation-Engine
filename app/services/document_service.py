"""
Document processing service for extracting text from various file formats.
"""
import logging
import io
from typing import Optional, Union
import PyPDF2
import docx
from fastapi import UploadFile

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for processing and extracting text from various document formats."""
    
    def __init__(self):
        """Initialize the document service."""
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
    
    async def extract_text_from_upload(self, file: UploadFile) -> str:
        """
        Extract text from an uploaded file.
        
        Args:
            file: Uploaded file object
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            Exception: If text extraction fails
        """
        if not file.filename:
            raise ValueError("File must have a filename")
        
        # Get file extension
        filename_lower = file.filename.lower()
        extension = None
        for ext in self.supported_extensions:
            if filename_lower.endswith(ext):
                extension = ext
                break
        
        if not extension:
            raise ValueError(
                f"Unsupported file type. Supported formats: {', '.join(self.supported_extensions)}"
            )
        
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        try:
            if extension == '.pdf':
                return self._extract_text_from_pdf(content)
            elif extension == '.docx':
                return self._extract_text_from_docx(content)
            elif extension == '.txt':
                return self._extract_text_from_txt(content)
            else:
                raise ValueError(f"Unsupported file extension: {extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file.filename}: {e}")
            raise Exception(f"Failed to extract text from {file.filename}: {str(e)}")
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF content.
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the PDF")
            
            logger.info(f"Extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """
        Extract text from DOCX content.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the DOCX file")
            
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_text_from_txt(self, content: bytes) -> str:
        """
        Extract text from TXT content.
        
        Args:
            content: TXT file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    extracted_text = content.decode(encoding)
                    if extracted_text.strip():
                        logger.info(f"Extracted {len(extracted_text)} characters from TXT (encoding: {encoding})")
                        return extracted_text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"TXT text extraction failed: {e}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    def validate_file_size(self, file_size: int, max_size: int) -> bool:
        """
        Validate if file size is within allowed limits.
        
        Args:
            file_size: Size of the file in bytes
            max_size: Maximum allowed size in bytes
            
        Returns:
            True if file size is valid, False otherwise
        """
        return file_size <= max_size
    
    def validate_file_extension(self, filename: str) -> bool:
        """
        Validate if file extension is supported.
        
        Args:
            filename: Name of the file
            
        Returns:
            True if extension is supported, False otherwise
        """
        if not filename:
            return False
        
        filename_lower = filename.lower()
        return any(filename_lower.endswith(ext) for ext in self.supported_extensions)
    
    def clean_extracted_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        cleaned = ' '.join(text.split())
        
        # Remove common document artifacts
        artifacts = [
            'Microsoft Word - ',
            'Page 1 of',
            'Page 2 of',
            'Page 3 of',
            'Page 4 of',
            'Page 5 of',
        ]
        
        for artifact in artifacts:
            cleaned = cleaned.replace(artifact, '')
        
        # Remove extra newlines and spaces
        cleaned = '\n'.join(line.strip() for line in cleaned.split('\n') if line.strip())
        
        return cleaned.strip()


# Global document service instance
document_service = DocumentService()
